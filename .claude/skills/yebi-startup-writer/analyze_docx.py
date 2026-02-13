from docx import Document
import sys

try:
    doc = Document('references/2025-yebi-startup.docx')

    print("=== 문서 구조 분석 ===\n")
    print(f"총 단락 수: {len(doc.paragraphs)}")
    print(f"총 표 수: {len(doc.tables)}\n")

    print("=== 첫 50개 단락 미리보기 ===\n")
    for i, para in enumerate(doc.paragraphs[:50]):
        if para.text.strip():
            style = para.style.name
            text = para.text[:100].replace('\n', ' ')
            print(f"{i:3d} [{style:25s}] {text}")

    print("\n=== 표 구조 ===\n")
    for i, table in enumerate(doc.tables[:10]):
        print(f"\n표 {i}: {len(table.rows)}행 × {len(table.columns)}열")
        if table.rows:
            print("첫 3행:")
            for row_idx in range(min(3, len(table.rows))):
                row_text = [cell.text[:40].replace('\n', ' ').strip() for cell in table.rows[row_idx].cells]
                print(f"  행{row_idx}: {row_text}")

    print("\n=== 플레이스홀더 검색 ===\n")
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if any(marker in text for marker in ['{{', '[]', '[입력]', '___', '예)', '(예:']):
            print(f"단락 {i}: {text[:80]}")

except Exception as e:
    print(f"Error: {e}", file=sys.stderr)
    sys.exit(1)
